from flask import Flask, request, jsonify
import openai
import docx

app = Flask(__name__)
openai.api_key = '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'  # Set your OpenAI API key

session_state = None

def extract_conversations_from_docx(file_path):
    doc = docx.Document(file_path)
    conversations = []
    current_conversation = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("User:"):
            current_conversation.append(paragraph.text.replace("User:", "").strip())
        elif paragraph.text.startswith("Bot:"):
            current_conversation.append(paragraph.text.replace("Bot:", "").strip())
            if idx < len(doc.paragraphs) - 1:
                next_paragraph = doc.paragraphs[idx + 1]
                if not next_paragraph.text.startswith(("User:", "Bot:")):
                    conversations.append(current_conversation)
                    current_conversation = []
            else:
                conversations.append(current_conversation)
    return conversations

def train_chatbot(conversations):
    global session_state
    training_data = []
    for conversation in conversations:
        for i in range(0, len(conversation), 2):
            training_data.append({"role": "system", "content": "You are a helpful assistant."})
            training_data.append({"role": "user", "content": conversation[i]})
            training_data.append({"role": "assistant", "content": conversation[i + 1]})
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=training_data,
        n=1,
        stop=None,
        temperature=0.7,
        max_tokens=100,
    )
    session_state = response['id']
    return session_state

def make_chatbot_request(session_id, messages):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        token=session_id,
        n=1,
        stop=None,
        temperature=0.7,
        max_tokens=100,
    )
    return response.choices[0].message['content']

@app.route('/train', methods=['POST'])
def train_chatbot_endpoint():
    global session_state
    session_state = None
    file = request.files['file']
    file.save('conversation.docx')  # Save the uploaded file
    conversations = extract_conversations_from_docx('conversation.docx')  # Extract conversations
    session_state = train_chatbot(conversations)  # Train the chatbot and get the session ID
    return jsonify({'session_id': session_state, 'message': 'Chatbot trained successfully!'})

@app.route('/chat', methods=['POST'])
def chat_with_bot_endpoint():
    global session_state
    token = request.headers.get('Authorization')
    session_id = request.json['session_id']
    messages = request.json['messages']
    
    if session_id:
        session_state = session_id
    
    response = make_chatbot_request(session_state, messages)  # Make request to ChatGPT
    return jsonify({'response': response})

if __name__ == '__main__':
    app.run()
